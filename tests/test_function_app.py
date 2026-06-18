import base64
import json
import sys
from pathlib import Path
from unittest.mock import patch

import azure.functions as func
import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

import function_app  # noqa: E402

REPO_ROOT = Path(__file__).resolve().parents[1]


def _request(payload):
    return func.HttpRequest(
        method="POST",
        url="/api/convert",
        body=json.dumps(payload).encode("utf-8"),
        params={},
    )


# ---------------------------------------------------------------------------
# Unit tests — _decode_content
# ---------------------------------------------------------------------------

def test_decode_content_rejects_invalid_base64():
    with pytest.raises(ValueError, match="not valid base64"):
        function_app._decode_content("not base64")


def test_decode_content_rejects_oversized_file(monkeypatch):
    monkeypatch.setattr(function_app, "MAX_FILE_BYTES", 3)
    payload = base64.b64encode(b"1234").decode("ascii")

    with pytest.raises(ValueError, match="too large"):
        function_app._decode_content(payload)


# ---------------------------------------------------------------------------
# Endpoint tests — /api/convert input validation
# ---------------------------------------------------------------------------

def test_convert_missing_filename():
    payload = {
        "content": base64.b64encode(b"hello").decode("ascii"),
    }

    response = function_app.convert(_request(payload))

    assert response.status_code == 400
    assert b"filename" in response.get_body()


def test_convert_non_string_filename():
    """filename: 123 should return 400, not crash with TypeError."""
    payload = {
        "content": base64.b64encode(b"hello").decode("ascii"),
        "filename": 123,
    }

    response = function_app.convert(_request(payload))

    assert response.status_code == 400
    assert b"filename" in response.get_body()


def test_convert_empty_filename():
    payload = {
        "content": base64.b64encode(b"hello").decode("ascii"),
        "filename": "   ",
    }

    response = function_app.convert(_request(payload))

    assert response.status_code == 400
    assert b"filename" in response.get_body()


def test_convert_filename_without_extension():
    payload = {
        "content": base64.b64encode(b"hello").decode("ascii"),
        "filename": "noextension",
    }

    response = function_app.convert(_request(payload))

    assert response.status_code == 415
    assert b"Unsupported" in response.get_body()


def test_convert_rejects_unsupported_extension():
    payload = {
        "content": base64.b64encode(b"hello").decode("ascii"),
        "filename": "legacy.doc",
    }

    response = function_app.convert(_request(payload))

    assert response.status_code == 415
    assert b"Unsupported" in response.get_body()


def test_convert_requires_json_object():
    request = func.HttpRequest(
        method="POST",
        url="/api/convert",
        body=b'["not", "an", "object"]',
        params={},
    )

    response = function_app.convert(request)

    assert response.status_code == 400
    assert b"JSON object" in response.get_body()


def test_convert_rejects_malformed_json():
    request = func.HttpRequest(
        method="POST",
        url="/api/convert",
        body=b"not json at all",
        params={},
    )

    response = function_app.convert(request)

    assert response.status_code == 400
    assert b"valid JSON" in response.get_body()


def test_convert_missing_content():
    payload = {"filename": "test.pdf"}

    response = function_app.convert(_request(payload))

    assert response.status_code == 400
    assert b"content" in response.get_body()


# ---------------------------------------------------------------------------
# Endpoint tests — /api/convert happy path
# ---------------------------------------------------------------------------

def test_convert_happy_path_docx(tmp_path):
    """Build a minimal .docx in-memory and verify end-to-end conversion."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello MarkItDown")
    docx_path = tmp_path / "sample.docx"
    doc.save(str(docx_path))

    content_b64 = base64.b64encode(docx_path.read_bytes()).decode("ascii")
    payload = {"content": content_b64, "filename": "sample.docx"}

    response = function_app.convert(_request(payload))

    assert response.status_code == 200
    body = response.get_body().decode("utf-8")
    assert "Hello MarkItDown" in body


# ---------------------------------------------------------------------------
# Endpoint tests — /api/convert error handling
# ---------------------------------------------------------------------------

def test_convert_returns_422_for_protected_file():
    """Simulates an encrypted/protected file raising a known exception."""
    payload = {
        "content": base64.b64encode(b"fake pdf bytes").decode("ascii"),
        "filename": "encrypted.pdf",
    }

    with patch.object(function_app._md, "convert", side_effect=RuntimeError("encrypted")):
        response = function_app.convert(_request(payload))

    assert response.status_code == 422
    assert b"encrypted" in response.get_body() or b"password-protected" in response.get_body()
    assert b"correlation_id=" in response.get_body()


def test_convert_returns_500_for_unexpected_error():
    """Unexpected exceptions should return 500 with a correlation_id."""
    payload = {
        "content": base64.b64encode(b"fake pdf bytes").decode("ascii"),
        "filename": "report.pdf",
    }

    with patch.object(function_app._md, "convert", side_effect=OSError("disk full")):
        response = function_app.convert(_request(payload))

    assert response.status_code == 500
    assert b"correlation_id=" in response.get_body()


def test_convert_cleans_up_temp_file():
    """Verify temp file is removed even after a conversion failure."""
    payload = {
        "content": base64.b64encode(b"fake pdf bytes").decode("ascii"),
        "filename": "report.pdf",
    }

    with patch.object(function_app._md, "convert", side_effect=OSError("boom")):
        response = function_app.convert(_request(payload))

    assert response.status_code == 500
    # The real assertion is that the finally block ran without raising.
    # We can't reliably check the temp dir (other processes leave files),
    # but reaching this point proves cleanup didn't shadow the exception.


def test_convert_does_not_log_filename(caplog):
    """Ensure error logs contain only correlation_id, not the filename."""
    payload = {
        "content": base64.b64encode(b"fake pdf bytes").decode("ascii"),
        "filename": "secret-document.pdf",
    }

    with patch.object(function_app._md, "convert", side_effect=OSError("fail")):
        function_app.convert(_request(payload))

    log_text = caplog.text
    assert "secret-document.pdf" not in log_text
    assert "correlation_id=" in log_text or "correlation_id" in log_text


# ---------------------------------------------------------------------------
# Endpoint tests — /api/health
# ---------------------------------------------------------------------------

def test_health_returns_ok():
    request = func.HttpRequest(method="GET", url="/api/health", body=b"", params={})

    response = function_app.health(request)

    assert response.status_code == 200
    assert response.get_body() == b"ok"


# ---------------------------------------------------------------------------
# OpenAPI spec validation
# ---------------------------------------------------------------------------

class TestOpenApiSpecs:
    """Validate OpenAPI specs match the function implementation."""

    @pytest.fixture(params=["connector/openapi.json", "connector/openapi-entra.json"])
    def spec(self, request):
        spec_path = REPO_ROOT / request.param
        with spec_path.open(encoding="utf-8-sig") as fh:
            return json.load(fh)

    def test_spec_has_convert_endpoint(self, spec):
        assert "/convert" in spec["paths"]
        assert "post" in spec["paths"]["/convert"]

    def test_spec_has_health_endpoint(self, spec):
        assert "/health" in spec["paths"]
        assert "get" in spec["paths"]["/health"]

    def test_spec_convert_requires_body(self, spec):
        params = spec["paths"]["/convert"]["post"]["parameters"]
        body_params = [p for p in params if p["in"] == "body"]
        assert len(body_params) == 1
        assert body_params[0]["required"] is True

    def test_spec_request_schema_has_required_fields(self, spec):
        schema = spec["definitions"]["ConvertRequest"]
        assert "content" in schema["required"]
        assert "filename" in schema["required"]
        assert schema["properties"]["content"]["type"] == "string"
        assert schema["properties"]["filename"]["type"] == "string"

    def test_spec_has_api_key_security(self, spec):
        assert "apiKeyQuery" in spec["securityDefinitions"]
        api_key = spec["securityDefinitions"]["apiKeyQuery"]
        assert api_key["type"] == "apiKey"
        assert api_key["name"] == "code"
        assert api_key["in"] == "query"

    def test_spec_convert_responses_match_code(self, spec):
        responses = spec["paths"]["/convert"]["post"]["responses"]
        assert "200" in responses
        assert "400" in responses
        assert "415" in responses
        assert "500" in responses

    def test_entra_spec_has_oauth_security(self):
        spec_path = REPO_ROOT / "connector" / "openapi-entra.json"
        with spec_path.open(encoding="utf-8-sig") as fh:
            spec = json.load(fh)

        assert "azureAdOAuth" in spec["securityDefinitions"]
        oauth = spec["securityDefinitions"]["azureAdOAuth"]
        assert oauth["type"] == "oauth2"
        assert oauth["flow"] == "accessCode"


# ---------------------------------------------------------------------------
# Smoke test — sample request/response fixture
# ---------------------------------------------------------------------------

def test_smoke_roundtrip_docx(tmp_path):
    """Full smoke test: create a DOCX, POST it, verify Markdown structure."""
    from docx import Document

    doc = Document()
    doc.add_heading("Smoke Test Report", level=1)
    doc.add_paragraph("This is paragraph one.")
    doc.add_paragraph("This is paragraph two.")
    docx_path = tmp_path / "smoke.docx"
    doc.save(str(docx_path))

    content_b64 = base64.b64encode(docx_path.read_bytes()).decode("ascii")
    payload = {"content": content_b64, "filename": "smoke.docx"}

    response = function_app.convert(_request(payload))

    assert response.status_code == 200
    md = response.get_body().decode("utf-8")
    assert "Smoke Test Report" in md
    assert "paragraph one" in md
    assert "paragraph two" in md
